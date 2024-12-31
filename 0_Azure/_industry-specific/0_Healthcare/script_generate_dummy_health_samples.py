# Create new samples with more than 1,000 rows for CSV files and around 15 pages for DOCX files

import pandas as pd
import numpy as np
from docx import Document

# Define the number of samples
num_samples = 1500

# 1. Medical History Records
medical_history_data = {
    'PatientID': np.arange(1, num_samples + 1),
    'Name': [f'Patient_{i}' for i in range(1, num_samples + 1)],
    'Age': np.random.randint(0, 100, num_samples),
    'Gender': np.random.choice(['Male', 'Female'], num_samples),
    'PastMedicalConditions': np.random.choice(['Hypertension', 'Diabetes', 'Asthma', 'None'], num_samples),
    'Surgeries': np.random.choice(['Appendectomy', 'None', 'Gallbladder removal'], num_samples),
    'Allergies': np.random.choice(['Penicillin', 'None', 'Peanuts'], num_samples),
    'FamilyMedicalHistory': np.random.choice(['Heart Disease', 'Diabetes', 'None'], num_samples)
}
medical_history_df = pd.DataFrame(medical_history_data)
medical_history_df.to_csv('large_medical_history_records.csv', index=False)

# 2. Lab Test Results
lab_test_results_data = {
    'PatientID': np.arange(1, num_samples + 1),
    'TestName': np.random.choice(['Complete Blood Count (CBC)', 'Liver Function Test', 'Kidney Function Test'], num_samples),
    'Result': np.random.choice(['Normal', 'Elevated ALT', 'Low Hemoglobin'], num_samples),
    'Date': pd.date_range(start='2023-01-01', periods=num_samples, freq='D')
}
lab_test_results_df = pd.DataFrame(lab_test_results_data)
lab_test_results_df.to_csv('large_lab_test_results.csv', index=False)

# 3. Radiology Reports
radiology_report = Document()
for i in range(1, 16):
    radiology_report.add_heading(f'Radiology Report - Page {i}', 0)
    radiology_report.add_paragraph(f'PatientID: {i}')
    radiology_report.add_paragraph(f'Name: Patient_{i}')
    radiology_report.add_paragraph('Imaging Study: Chest X-ray')
    radiology_report.add_paragraph('Findings: No acute cardiopulmonary process.')
    radiology_report.add_paragraph('Impression: Normal chest X-ray.')
radiology_report.save('large_radiology_report.docx')

# 4. Prescription Records
prescription_data = {
    'PatientID': np.arange(1, num_samples + 1),
    'Medication': np.random.choice(['Lisinopril', 'Albuterol', 'Metformin'], num_samples),
    'Dosage': np.random.choice(['10mg', '90mcg', '500mg'], num_samples),
    'Frequency': np.random.choice(['Once daily', 'As needed', 'Twice daily'], num_samples)
}
prescription_df = pd.DataFrame(prescription_data)
prescription_df.to_csv('large_prescription_records.csv', index=False)

# 5. Vital Signs Logs
vital_signs_data = {
    'PatientID': np.arange(1, num_samples + 1),
    'Date': pd.date_range(start='2023-01-01', periods=num_samples, freq='D'),
    'Temperature_C': np.random.normal(37, 0.5, num_samples).astype(float),
    'Pulse_bpm': np.random.normal(70, 10, num_samples).astype(int),
    'RespirationRate_bpm': np.random.normal(16, 2, num_samples).astype(int),
    'BloodPressure_systolic': np.random.normal(120, 15, num_samples).astype(int),
    'BloodPressure_diastolic': np.random.normal(80, 10, num_samples).astype(int)
}
vital_signs_df = pd.DataFrame(vital_signs_data)
vital_signs_df.to_csv('large_vital_signs_logs.csv', index=False)

# 6. Immunization Records
immunization_data = {
    'PatientID': np.arange(1, num_samples + 1),
    'Vaccine': np.random.choice(['Influenza', 'COVID-19', 'Hepatitis B'], num_samples),
    'DateAdministered': pd.date_range(start='2022-01-01', periods=num_samples, freq='D'),
    'DoseNumber': np.random.choice([1, 2, 3], num_samples)
}
immunization_df = pd.DataFrame(immunization_data)
immunization_df.to_csv('large_immunization_records.csv', index=False)

# 7. Clinical Notes
clinical_notes = Document()
for i in range(1, 16):
    clinical_notes.add_heading(f'Clinical Notes - Page {i}', 0)
    clinical_notes.add_paragraph(f'PatientID: {i}')
    clinical_notes.add_paragraph(f'Name: Patient_{i}')
    clinical_notes.add_paragraph('Visit Date: 2023-03-05')
    clinical_notes.add_paragraph('Chief Complaint: Cough and fever.')
    clinical_notes.add_paragraph('Assessment: Likely viral upper respiratory infection.')
    clinical_notes.add_paragraph('Plan: Symptomatic treatment and follow-up in one week if no improvement.')
clinical_notes.save('large_clinical_notes.docx')

# 8. Surgical Reports
surgical_report = Document()
for i in range(1, 16):
    surgical_report.add_heading(f'Surgical Report - Page {i}', 0)
    surgical_report.add_paragraph(f'PatientID: {i}')
    surgical_report.add_paragraph(f'Name: Patient_{i}')
    surgical_report.add_paragraph('Procedure: Appendectomy')
    surgical_report.add_paragraph('Date of Surgery: 2023-02-25')
    surgical_report.add_paragraph('Pre-operative Diagnosis: Acute appendicitis.')
    surgical_report.add_paragraph('Post-operative Diagnosis: Acute appendicitis.')
    surgical_report.add_paragraph('Procedure Details: Laparoscopic appendectomy performed without complications.')
surgical_report.save('large_surgical_report.docx')

# 9. Mental Health Records
mental_health_data = {
    'PatientID': np.arange(1, num_samples + 1),
    'AssessmentDate': pd.date_range(start='2023-01-01', periods=num_samples, freq='D'),
    'Diagnosis': np.random.choice(['Generalized Anxiety Disorder', 'Major Depressive Disorder'], num_samples),
    'TreatmentPlan': np.random.choice(['Cognitive Behavioral Therapy (CBT)', 
                                       'Selective Serotonin Reuptake Inhibitor (SSRI)'], 
                                      num_samples)
}
mental_health_df = pd.DataFrame(mental_health_data)
mental_health_df.to_csv('large_mental_health_records.csv', index=False)

# 10. Diet and Nutrition Logs
diet_nutrition_data = {
    'PatientID': np.arange(1, num_samples + 1),
    'Date': pd.date_range(start='2023-01-01', periods=num_samples, freq='D'),
    'MealType': np.random.choice(['Breakfast', 'Lunch', 'Dinner'], num_samples),
    'FoodItems': np.random.choice(['Oatmeal, Banana, Orange Juice',
                                   'Grilled Chicken Salad, Apple',
                                   'Pasta, Tomato Sauce, Salad'], num_samples)
}
diet_nutrition_df = pd.DataFrame(diet_nutrition_data)
diet_nutrition_df.to_csv('large_diet_nutrition_logs.csv', index=False)

print("Larger sample health documents have been created and saved in appropriate formats.")
